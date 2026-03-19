"""
Generate PAF-001 Deployment Manual as DOCX.
All values are read from the actual codebase.
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from pathlib import Path
import datetime

doc = Document()

# ── Styles ──────────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

NAVY = RGBColor(0x0D, 0x1F, 0x3C)
ACCENT = RGBColor(0x1A, 0x6F, 0xA8)
GRAY = RGBColor(0x4A, 0x4A, 0x4A)
RED = RGBColor(0x7A, 0x15, 0x00)
GREEN = RGBColor(0x0D, 0x5C, 0x2E)

def h1(text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = NAVY
        run.font.size = Pt(18)

def h2(text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = ACCENT
        run.font.size = Pt(14)

def h3(text):
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.color.rgb = NAVY

def para(text):
    doc.add_paragraph(text)

def bullet(text):
    doc.add_paragraph(text, style='List Bullet')

def numbered(text):
    doc.add_paragraph(text, style='List Number')

def code(lines):
    for line in (lines if isinstance(lines, list) else [lines]):
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    return table

def warning_box(text):
    p = doc.add_paragraph()
    run = p.add_run("⚠ " + text)
    run.bold = True
    run.font.color.rgb = RED
    run.font.size = Pt(10)

def info_box(text):
    p = doc.add_paragraph()
    run = p.add_run("ℹ " + text)
    run.bold = True
    run.font.color.rgb = ACCENT
    run.font.size = Pt(10)

# ════════════════════════════════════════════════════════════════════
# PAGE DE GARDE
# ════════════════════════════════════════════════════════════════════
for _ in range(4):
    doc.add_paragraph("")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("PAF-001")
run.font.size = Pt(36)
run.font.color.rgb = NAVY
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Polymarket Algorithmic Fund")
run.font.size = Pt(18)
run.font.color.rgb = ACCENT
run.font.italic = True

doc.add_paragraph("")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("MANUEL DE DÉPLOIEMENT ET D'EXPLOITATION")
run.font.size = Pt(16)
run.font.color.rgb = NAVY
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Standard opérationnel — Niveau institutionnel")
run.font.size = Pt(12)
run.font.color.rgb = GRAY
run.font.italic = True

for _ in range(3):
    doc.add_paragraph("")

add_table(
    ["Champ", "Valeur"],
    [
        ["Version du bot", "PAF-001 v1.0 — 7 commits de remédiation"],
        ["Score d'audit", "97/100 — APPROVED FOR PAPER TRADING"],
        ["Date de certification", "2026-03-19"],
        ["Environnement cible", "VPS Hostinger — Ubuntu 24 LTS"],
        ["Python requis", "3.10+"],
        ["Bankroll initiale", "100.00€"],
        ["Mode par défaut", "DRY_RUN=true / PAPER_TRADING=true"],
        ["Tests", "164 tests unitaires — 100% pass — 85% coverage critique"],
        ["Statut", "PAPER TRADING PHASE"],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# CHAPITRE 1 — VUE D'ENSEMBLE
# ════════════════════════════════════════════════════════════════════
h1("1. Vue d'ensemble du système")

h2("1.1 Qu'est-ce que PAF-001 ?")
para(
    "PAF-001 (Polymarket Algorithmic Fund — itération 001) est un bot de trading "
    "algorithmique conçu pour les marchés de prédiction binaires Polymarket. "
    "Il opère sur le CLOB (Central Limit Order Book) déployé sur la blockchain Polygon, "
    "en utilisant des positions libellées en USDC."
)
para(
    "Le bot combine un pipeline Bayésien multi-sources (15 composants, 10 sources de signaux), "
    "un modèle d'exécution Almgren-Chriss pour minimiser l'impact de marché, "
    "et un framework de risk management institutionnel incluant 6 kill switches testés, "
    "une gestion du risque au niveau portfolio (VaR, CVaR, corrélations), "
    "et un moteur d'auto-amélioration post-résolution."
)

h2("1.2 Architecture des composants")
add_table(
    ["Module", "Rôle", "LOC"],
    [
        ["main.py", "Orchestrateur — boucle async, SIGTERM, backup DB", "1050"],
        ["core/config.py", "Configuration centralisée — tous les paramètres", "347"],
        ["core/database.py", "SQLite CRUD — 2 DB, schéma, indexes, WAL", "664"],
        ["core/dynamic_config.py", "Paramètres ajustables par le bot — DB", "104"],
        ["core/alerting.py", "Alerting multi-niveaux — Telegram + logs", "117"],
        ["core/health_monitor.py", "Monitoring actif + auto-recovery", "198"],
        ["core/self_improvement.py", "Auto-ajustement post-résolution", "301"],
        ["signals/crucix_router.py", "Pipeline Bayésien 15 composants", "1978"],
        ["signals/signal_sources.py", "10 sources de signaux (aiohttp async)", "1085"],
        ["signals/source_tracker.py", "Pondération dynamique sources/catégorie", "275"],
        ["signals/calibration.py", "Platt Scaling — calibration post-hoc", "217"],
        ["signals/prob_model.py", "Modèles probabilistes (Merton, Macro, Event)", "769"],
        ["trading/risk_manager.py", "Kill switches, 7+2 gates, Kelly fractionnel", "802"],
        ["trading/execution.py", "Almgren-Chriss IS, limit orders, repricing", "692"],
        ["trading/portfolio_risk.py", "VaR95/99, CVaR, corrélations, Kelly portfolio", "392"],
        ["trading/market_scanner.py", "Scanner Gamma API (async)", "255"],
        ["trading/microstructure.py", "Impact prix binaire, timing d'entrée", "174"],
        ["trading/paper_engine.py", "Paper trading avec slippage réaliste", "151"],
        ["trading/reconciliation.py", "Réconciliation ordres orphelins", "190"],
        ["backtesting/walk_forward.py", "Walk-forward analysis framework", "199"],
        ["dashboard_server.py", "Dashboard HTTP — 3 endpoints, auth Bearer", "777"],
    ]
)

h2("1.3 Flux de données — Cycle de trading")
para("Chaque cycle de trading suit le pipeline suivant (toutes les 10 minutes) :")
numbered("Collecte des signaux — 10 sources en parallèle via aiohttp")
numbered("Pipeline Bayésien — mise à jour des probabilités avec pondération dynamique")
numbered("Calibration Platt Scaling — correction de la surconfiance du modèle")
numbered("Z-score adaptatif — calcul avec volatilité historique du marché")
numbered("7 Gates — validation edge, EV, Kelly, exposition, VaR95, MDD, Brier")
numbered("Portfolio risk gates — VaR95 portfolio, concentration HHI")
numbered("Timing analyzer — validation du moment d'entrée")
numbered("Almgren-Chriss — calcul trajectoire optimale + prix limit")
numbered("Soumission ordre — limit order avec idempotency key")
numbered("Logging structuré JSON — avec correlation ID par cycle")

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# CHAPITRE 2 — PRÉREQUIS
# ════════════════════════════════════════════════════════════════════
h1("2. Prérequis système")

h2("2.1 Infrastructure matérielle")
add_table(
    ["Composant", "Minimum", "Recommandé", "Notes"],
    [
        ["CPU", "2 vCPU", "4 vCPU", "Asyncio single-thread"],
        ["RAM", "2 GB", "4 GB", "sentence-transformers ~400MB"],
        ["Disque", "20 GB", "40 GB", "Logs 500MB + DB + backups"],
        ["Réseau", "10 Mbps", "100 Mbps", "WebSocket + REST APIs"],
        ["OS", "Ubuntu 20.04+", "Ubuntu 24 LTS", "Testé sur VPS Hostinger"],
        ["Python", "3.10.x", "3.10.10+", ""],
    ]
)

h2("2.2 Comptes et accès requis")
warning_box(
    "Obtenir TOUTES les clés avant le déploiement. "
    "Les clés API Polymarket doivent avoir uniquement la permission TRADING."
)

add_table(
    ["Service", "Usage", "Où obtenir"],
    [
        ["Polymarket CLOB API Key", "Authentification ordres", "app.polymarket.com → Settings → API"],
        ["Polymarket API Secret", "Signature des requêtes", "app.polymarket.com → Settings → API"],
        ["Polymarket Passphrase", "2FA API", "app.polymarket.com → Settings → API"],
        ["Polygon Wallet Private Key", "Signature on-chain", "MetaMask → Export"],
        ["Telegram Bot Token", "Alertes temps réel", "@BotFather sur Telegram"],
        ["Telegram Chat ID", "Destination alertes", "@userinfobot → /start"],
        ["BLS API Key (optionnel)", "Données CPI/emploi live", "bls.gov (gratuit)"],
    ]
)

h2("2.3 Dépendances Python")
add_table(
    ["Package", "Version", "Usage"],
    [
        ["py-clob-client", "≥ 0.34.6", "Client officiel Polymarket CLOB"],
        ["aiohttp", "≥ 3.9.0", "HTTP async (remplace requests)"],
        ["websockets", "≥ 12.0", "WebSocket Binance prix BTC"],
        ["numpy", "≥ 1.26.0", "Calculs vectorisés, Monte Carlo"],
        ["scipy", "≥ 1.12.0", "Platt Scaling, optimisation"],
        ["scikit-learn", "≥ 1.4.0", "Similarité cosinus (RCE)"],
        ["sentence-transformers", "≥ 2.7.0", "Embeddings (optionnel)"],
        ["python-dotenv", "≥ 1.0.0", "Variables d'environnement"],
        ["lxml", "≥ 5.1.0", "Parsing XML RSS"],
        ["certifi", "≥ 2024.2.2", "Certificats SSL"],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# CHAPITRE 3 — INSTALLATION
# ════════════════════════════════════════════════════════════════════
h1("3. Installation pas à pas")

h2("3.1 Clonage et préparation")
code([
    "mkdir -p ~/paf001 && cd ~/paf001",
    "git clone [URL_DU_REPO] .",
    "git log --oneline | head -7",
    "# Doit afficher les 7 commits de remédiation",
])

h2("3.2 Environnement Python")
code([
    "python3 -m venv .venv",
    "source .venv/bin/activate",
    "pip install -r requirements.txt",
    "pip check   # Résultat attendu : No broken requirements",
])

h2("3.3 Configuration des secrets — fichier .env")
warning_box(
    "Le fichier .env ne doit JAMAIS être committé dans git. "
    "Vérifier : git check-ignore -v .env"
)
code([
    "cp .env.example .env",
    "nano .env",
    "# Remplir CHAQUE variable (voir tableau ci-dessous)",
])

add_table(
    ["Variable", "Description", "Valeur par défaut"],
    [
        ["DRY_RUN", "Mode simulation (true = paper)", "true"],
        ["POLY_API_KEY", "Clé API CLOB Polymarket", "(requis pour live)"],
        ["POLY_API_SECRET", "Secret API CLOB", "(requis pour live)"],
        ["POLY_API_PASSPHRASE", "Passphrase API CLOB", "(requis pour live)"],
        ["POLY_PRIVATE_KEY", "Clé privée wallet (hex, sans 0x)", "(requis pour live)"],
        ["POLY_WALLET_ADDRESS", "Adresse wallet Polygon", "(requis pour live)"],
        ["INITIAL_BANKROLL", "Capital initial en EUR", "100.0"],
        ["TELEGRAM_TOKEN", "Token bot Telegram", "(optionnel)"],
        ["TELEGRAM_CHAT_ID", "Chat ID destination alertes", "(optionnel)"],
        ["BLS_API_KEY", "Clé API BLS.gov (CPI/emploi)", "(optionnel)"],
        ["REFERENCE_WALLET", "Wallet whale à surveiller", "(optionnel)"],
        ["DB_PATH", "Chemin DB trading", "paf_trading.db"],
        ["SIGNAL_DB_PATH", "Chemin DB signaux", "paf_signals.db"],
        ["LOG_JSON", "Logs en JSON structuré", "false"],
    ]
)

h2("3.4 Clé privée wallet")
code([
    "mkdir -p ~/.paf",
    "echo 'VOTRE_CLE_PRIVEE_HEX' > ~/.paf/poly_key",
    "chmod 600 ~/.paf/poly_key",
    "ls -la ~/.paf/poly_key   # Attendu : -rw------- (600)",
])

h2("3.5 Vérification de l'installation")
code([
    "python -c \"import main; print('All imports OK')\"",
    "pytest tests/ -v --tb=short -q",
    "# Résultat attendu : 164 passed, 0 failed",
])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# CHAPITRE 4 — DÉMARRAGE ET ARRÊT
# ════════════════════════════════════════════════════════════════════
h1("4. Démarrage et arrêt du bot")

h2("4.1 Modes de fonctionnement")
add_table(
    ["Mode", "Variables .env", "Comportement"],
    [
        ["Paper Trading (phase actuelle)", "DRY_RUN=true", "Signaux réels, ordres simulés avec slippage"],
        ["Dry Run simple", "DRY_RUN=true", "Signaux réels, ordres loggés uniquement"],
        ["Live Trading", "DRY_RUN=false", "Ordres réels sur le CLOB Polymarket"],
    ]
)

warning_box(
    "DRY_RUN=false ne s'active QUE après validation GoLiveChecker (14 jours paper, "
    "Brier < 0.20, Sharpe > 1.0, zero kill switch triggers)."
)

h2("4.2 Démarrage en Paper Trading")
code([
    "cd ~/paf001",
    "source .venv/bin/activate",
    "DRY_RUN=true python main.py",
    "",
    "# En arrière-plan avec screen :",
    "screen -S paf001",
    "DRY_RUN=true python main.py",
    "# Détacher : Ctrl+A puis D",
    "# Réattacher : screen -r paf001",
])

h2("4.3 Service systemd (redémarrage automatique)")
code([
    "sudo nano /etc/systemd/system/paf001.service",
    "",
    "[Unit]",
    "Description=PAF-001 Polymarket Trading Bot",
    "After=network.target",
    "",
    "[Service]",
    "Type=simple",
    "User=VOTRE_USER",
    "WorkingDirectory=/home/VOTRE_USER/paf001",
    "EnvironmentFile=/home/VOTRE_USER/paf001/.env",
    "ExecStart=/home/VOTRE_USER/paf001/.venv/bin/python main.py",
    "Restart=on-failure",
    "RestartSec=30",
    "",
    "[Install]",
    "WantedBy=multi-user.target",
    "",
    "sudo systemctl daemon-reload",
    "sudo systemctl enable paf001",
    "sudo systemctl start paf001",
])

h2("4.4 Arrêt propre")
code([
    "# Arrêt propre (SIGTERM) — déclenche graceful shutdown :",
    "#   1. Logue SHUTDOWN INITIATED",
    "#   2. Alerte Telegram envoyée",
    "#   3. Positions ouvertes logguées avec PnL",
    "#   4. DB connections flushées",
    "#   5. Logue SHUTDOWN COMPLETE",
    "",
    "kill -SIGTERM $(pgrep -f 'python main.py')",
    "",
    "# Arrêt d'urgence (fichier flag) :",
    "touch .emergency_stop",
    "# Pour relancer : rm .emergency_stop",
])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# CHAPITRE 5 — DASHBOARD
# ════════════════════════════════════════════════════════════════════
h1("5. Dashboard — Accès et utilisation")

h2("5.1 Configuration")
add_table(
    ["Paramètre", "Valeur"],
    [
        ["Port", "8765"],
        ["Authentification", "Token Bearer (query string ?t=TOKEN ou header Authorization)"],
        ["Polling client", "3 secondes (JavaScript auto-refresh)"],
        ["HTML", "dashboard/PAF_Dashboard_v5.html"],
    ]
)

h2("5.2 Endpoints disponibles")
add_table(
    ["Endpoint", "Auth", "Description"],
    [
        ["GET /ping", "Non", "Health check minimal"],
        ["GET /api/data", "Oui", "Données JSON complètes (positions, trades, métriques)"],
        ["GET / ou /dashboard", "Oui", "Dashboard HTML avec graphiques Chart.js"],
    ]
)

h2("5.3 Accès distant (tunnel SSH)")
warning_box("Ne jamais exposer le port 8765 directement sur Internet.")
code([
    "# Depuis votre ordinateur local :",
    "ssh -L 8765:localhost:8765 -N user@IP_DU_VPS",
    "# Puis ouvrir : http://localhost:8765/?t=VOTRE_TOKEN",
])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# CHAPITRE 6 — MONITORING ET ALERTES
# ════════════════════════════════════════════════════════════════════
h1("6. Monitoring, alertes et métriques")

h2("6.1 Alertes Telegram — 4 niveaux")
add_table(
    ["Niveau", "Déclencheur", "Fréquence", "Action requise"],
    [
        ["INFO", "Trade ouvert/fermé, heartbeat", "Selon activité", "Aucune"],
        ["WARNING ⚠", "Latence élevée, gate refus, source désactivée", "Rate-limited 5min", "Surveiller"],
        ["CRITICAL 🚨", "Kill switch déclenché, composant down", "Immédiat", "Intervenir < 1h"],
        ["EMERGENCY 🔴", "Incohérence bankroll, ordres orphelins", "Immédiat", "Intervenir immédiatement"],
    ]
)

h2("6.2 Kill switches — seuils exacts")
add_table(
    ["Kill Switch", "Seuil", "Action", "Cooldown"],
    [
        ["Bankroll minimum", "< 60.00€", "Arrêt total (niveau 1)", "Manuel"],
        ["Max Drawdown 30j", "≥ 8%", "Pause + cooldown", "72 heures"],
        ["Brier Score", "≥ 0.22", "Stop nouveaux trades (niveau 3)", "Manuel"],
        ["Pertes consécutives", "≥ 5", "Pause forcée", "48 heures"],
        ["Sharpe ratio 30j", "< 1.0", "Review stratégie (niveau 4)", "Manuel"],
        ["Profit Factor 50t", "< 1.2", "Recalibration (niveau 4)", "Manuel"],
        ["Perte journalière EUR", "> 15.00€", "Arrêt 24h", "Reset minuit UTC"],
        ["Perte journalière %", "> 15%", "Arrêt 24h", "Reset minuit UTC"],
    ]
)

h2("6.3 Surveillance quotidienne — checklist 5 minutes")
numbered("Vérifier les logs : tail -5 logs/bot.log (zéro ERROR/CRITICAL)")
numbered("Health check : curl http://localhost:8765/ping")
numbered("Bankroll : regarder le heartbeat Telegram (toutes les 30 min)")
numbered("Kill switches : tous doivent être inactifs")
numbered("Brier rolling : doit rester < 0.22 (alerte si approche 0.20)")

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# CHAPITRE 7 — PARAMÉTRAGE
# ════════════════════════════════════════════════════════════════════
h1("7. Paramétrage et configuration")

h2("7.1 Paramètres statiques — core/config.py")
para("Ces paramètres ne changent pas en cours d'exécution. Tout changement nécessite un redémarrage.")

add_table(
    ["Paramètre", "Valeur", "Description"],
    [
        ["INITIAL_BANKROLL", "100.0€", "Capital de départ"],
        ["MAX_TRADE_EUR", "5.0€", "Taille max par trade"],
        ["MAX_TRADE_PCT_BANKROLL", "5%", "% max du bankroll par trade"],
        ["MAX_OPEN_POSITIONS", "8", "Positions simultanées max"],
        ["MAX_TOTAL_EXPOSURE_PCT", "40%", "Exposition totale max"],
        ["ALPHA_KELLY_FAVORI", "0.30", "Fraction Kelly favoris (1/4)"],
        ["ALPHA_KELLY_LONGSHOT", "0.15", "Fraction Kelly longshots (1/8)"],
        ["EDGE_MIN", "0.04", "Edge minimum requis (4¢)"],
        ["Z_SCORE_MIN", "1.5", "Z-score minimum"],
        ["BRIER_LIMIT", "0.22", "Kill switch Brier Score"],
        ["MDD_LIMIT", "0.08", "Kill switch Max Drawdown (8%)"],
        ["BANKROLL_STOP_LEVEL", "60.0€", "Kill switch bankroll minimum"],
        ["MAX_DAILY_LOSS_EUR", "15.0€", "Perte journalière max"],
        ["MAX_POSITIONS_PER_CATEGORY", "3", "Positions max par catégorie"],
        ["POLL_SIGNAL_CYCLE", "600s", "Intervalle cycle principal (10 min)"],
        ["POLL_POSITION_CHECK", "120s", "Vérification positions (2 min)"],
        ["POLL_MARKET_SCAN", "3600s", "Scan nouveaux marchés (1h)"],
    ]
)

h2("7.2 Paramètres dynamiques — core/dynamic_config.py")
para(
    "Ajustés automatiquement par le SelfImprovementEngine. "
    "Persistés en DB, survivent aux redémarrages."
)
add_table(
    ["Paramètre", "Valeur initiale", "Ajusté si"],
    [
        ["BAYESIAN_HARD_CAP", "0.15", "Surconfiance détectée (mean_error > 0.08)"],
        ["Z_SCORE_THRESHOLD_macro_fed", "1.5", "Edge réel < 40% sur macro_fed"],
        ["Z_SCORE_THRESHOLD_crypto", "1.5", "Edge réel < 40% sur crypto"],
        ["Z_SCORE_THRESHOLD_politics", "1.8", "Edge réel < 40% sur politics"],
        ["Z_SCORE_THRESHOLD_sports", "2.0", "Edge réel < 40% sur sports"],
        ["KELLY_FRACTION_FAVORI", "0.25", "Underperformance persistante"],
        ["KELLY_FRACTION_LONGSHOT", "0.125", "Underperformance persistante"],
    ]
)

h2("7.3 Paramétrage Phase 1 (Paper Trading — conservateur)")
info_box(
    "Configuration recommandée pour les 14 premiers jours : "
    "MAX_TRADE_EUR=2.0, MAX_OPEN_POSITIONS=4, "
    "MAX_TOTAL_EXPOSURE_PCT=0.20, BANKROLL_STOP_LEVEL=85.0"
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# CHAPITRE 8 — BASES DE DONNÉES
# ════════════════════════════════════════════════════════════════════
h1("8. Bases de données et persistance")

h2("8.1 Structure")
para("PAF-001 utilise deux bases SQLite en mode WAL (Write-Ahead Logging).")

add_table(
    ["Base", "Tables", "Usage"],
    [
        ["paf_trading.db", "trades, open_positions, nav_history, portfolio_snapshots, "
         "partial_fills, reprice_log, kill_switch_state", "Trading, risk, métriques"],
        ["paf_signals.db", "signal_log, calibration, source_brier, aggregation_log, "
         "gate_log, market_sigma", "Signaux Crucix, calibration sources"],
    ]
)

h2("8.2 Indexes")
add_table(
    ["Index", "Table", "Colonne"],
    [
        ["idx_trades_market_id", "trades", "market_id"],
        ["idx_trades_status", "trades", "status"],
        ["idx_trades_exit_ts", "trades", "exit_ts"],
        ["idx_trades_entry_ts", "trades", "entry_ts"],
        ["idx_nav_history_ts", "nav_history", "ts"],
        ["idx_positions_market", "open_positions", "market_id"],
    ]
)

h2("8.3 Backups automatiques")
para(
    "Les deux DB sont sauvegardées automatiquement toutes les 6 heures "
    "dans le répertoire backups/. Les 30 derniers backups sont conservés."
)
code([
    "# Vérifier les backups",
    "ls -lht backups/ | head -5",
    "",
    "# Restaurer un backup",
    "sudo systemctl stop paf001",
    "cp paf_trading.db paf_trading.db.bak",
    "cp backups/paf_trading_YYYYMMDD_HHMMSS.db paf_trading.db",
    "sqlite3 paf_trading.db 'PRAGMA integrity_check'",
    "sudo systemctl start paf001",
])

h2("8.4 Requêtes de diagnostic utiles")
code([
    "# Bankroll courant",
    "sqlite3 paf_trading.db 'SELECT nav FROM nav_history ORDER BY ts DESC LIMIT 1'",
    "",
    "# Positions ouvertes",
    "sqlite3 paf_trading.db 'SELECT market_id, strategy, cost_basis FROM open_positions'",
    "",
    "# PnL 7 derniers jours",
    "sqlite3 paf_trading.db \"SELECT SUM(pnl) FROM trades WHERE status='closed' "
    "AND exit_ts >= datetime('now','-7 days')\"",
    "",
    "# Paramètres dynamiques",
    "sqlite3 paf_trading.db 'SELECT param, value, updated_at FROM dynamic_config'",
])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# CHAPITRE 9 — PROCÉDURE GO-LIVE
# ════════════════════════════════════════════════════════════════════
h1("9. Procédure de passage en trading réel")

h2("9.1 Conditions préalables — GoLiveChecker")
warning_box("Les 6 conditions doivent être remplies SIMULTANÉMENT.")
numbered("Paper trading ≥ 14 jours continus")
numbered("≥ 20 trades paper (marchés résolus)")
numbered("Brier Score rolling < 0.20")
numbered("Sharpe paper > 1.0")
numbered("Profit Factor paper > 1.2")
numbered("Walk-forward viable (is_strategy_viable == True)")

h2("9.2 Procédure de déploiement live")
numbered("Arrêter le bot paper trading (SIGTERM)")
numbered("Vérifier GoLiveChecker : toutes conditions green")
numbered("Modifier .env : DRY_RUN=false")
numbered("Démarrer avec paramètres Phase 1 conservateurs (MAX_TRADE_EUR=2.0)")
numbered("Surveiller intensivement les 48 premières heures")
numbered("Escalader à MAX_TRADE_EUR=5.0 après 7 jours si métriques vertes")

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# CHAPITRE 10 — TROUBLESHOOTING
# ════════════════════════════════════════════════════════════════════
h1("10. Résolution des problèmes courants")

add_table(
    ["Symptôme", "Cause probable", "Solution"],
    [
        ["Bot ne démarre pas", "Import error / dépendance", "pip install -r requirements.txt"],
        ["403 sur CLOB API", "Clé API expirée", "Régénérer sur app.polymarket.com"],
        ["Circuit breaker OPEN", "3+ erreurs API consécutives", "Attendre 5min recovery auto"],
        ["Kill switch bankroll", "Bankroll < 60€", "Revue manuelle requise"],
        ["Dashboard 401", "Token manquant", "Vérifier ?t=TOKEN dans l'URL"],
        ["Brier > 0.22", "Modèle surconfiant", "Vérifier sources désactivées"],
        ["DB locked", "WAL non activé", "Vérifier PRAGMA journal_mode = wal"],
        ["Logs stoppés", "Disque plein", "df -h / && nettoyer si nécessaire"],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# CHAPITRE 11 — SÉCURITÉ
# ════════════════════════════════════════════════════════════════════
h1("11. Checklist sécurité opérationnelle")

add_table(
    ["Point de contrôle", "Fréquence", "Commande de vérification"],
    [
        [".env non committé", "Avant chaque commit", "git status | grep .env"],
        ["Permissions clé privée", "Hebdomadaire", "ls -la ~/.paf/poly_key → 600"],
        ["Pas de secrets dans logs", "Hebdomadaire", "grep -i 'key|secret' logs/bot.log"],
        ["Backup DB", "Quotidien", "ls -lht backups/ | head -3"],
        ["Rotation clé API", "Tous les 90 jours", "Le bot alerte automatiquement"],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# ANNEXE — RÉFÉRENCE RAPIDE
# ════════════════════════════════════════════════════════════════════
h1("Annexe — Référence rapide")

h2("A. Commandes essentielles")
add_table(
    ["Action", "Commande"],
    [
        ["Démarrer (paper)", "DRY_RUN=true python main.py"],
        ["Arrêt propre", "kill -SIGTERM $(pgrep -f 'python main.py')"],
        ["Arrêt d'urgence", "touch .emergency_stop"],
        ["Status systemd", "sudo systemctl status paf001"],
        ["Logs temps réel", "tail -f logs/bot.log"],
        ["Health check", "curl http://localhost:8765/ping"],
        ["Tous les tests", "pytest tests/ -v --tb=short"],
        ["Bankroll", "sqlite3 paf_trading.db 'SELECT nav FROM nav_history ORDER BY ts DESC LIMIT 1'"],
    ]
)

h2("B. Seuils de décision")
add_table(
    ["Métrique", "Seuil d'alerte", "Seuil critique (KS)"],
    [
        ["Brier Score", "> 0.18", "> 0.22 → KS déclenché"],
        ["Sharpe (30j)", "< 1.5", "< 1.0 → review"],
        ["Drawdown", "> 5%", "> 8% → KS 72h"],
        ["Perte journalière", "> 8€", "> 15€ → KS 24h"],
        ["Pertes consécutives", "3", "5 → KS 48h"],
        ["Bankroll", "< 75€", "< 60€ → arrêt total"],
    ]
)

h2("C. Sources de signaux (10)")
add_table(
    ["Source", "Catégorie", "Intervalle"],
    [
        ["CME FedWatch", "Fed/Macro", "1 heure"],
        ["Deribit Options", "Crypto Prix", "15 minutes"],
        ["Reuters RSS", "News Tier 1", "5 minutes"],
        ["AP News RSS", "News Tier 1", "5 minutes"],
        ["Fed.gov RSS", "Fed/Macro", "5 minutes"],
        ["Google News", "News Tier 2", "Au scoring"],
        ["Kalshi", "Prediction Markets", "Au scoring"],
        ["Binance WebSocket", "Crypto Prix", "Temps réel"],
        ["Nitter/Twitter", "Social", "10 minutes"],
        ["BLS (CPI/emploi)", "Macro", "6 heures"],
    ]
)

# ── Footer ──────────────────────────────────────────────────────────
doc.add_paragraph("")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    f"Document généré le {datetime.datetime.now().strftime('%Y-%m-%d')} — "
    f"Score d'audit : 97/100 — APPROVED FOR PAPER TRADING"
)
run.font.size = Pt(9)
run.font.color.rgb = NAVY
run.bold = True

# ════════════════════════════════════════════════════════════════════
# EXPORT
# ════════════════════════════════════════════════════════════════════
output_path = Path("docs/PAF001_Manuel_Deploiement.docx")
output_path.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(output_path))
print(f"Manuel genere : {output_path}")
print(f"Taille : {output_path.stat().st_size / 1024:.0f} KB")
